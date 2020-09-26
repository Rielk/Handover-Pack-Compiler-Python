# -*- coding: utf-8 -*-
"""
Created on Tue Aug 25 14:16:28 2020

The Handover Pack class for the compilation assiter

@author: William
"""
import backend
import traceback
import shutil
import find
import json
import ui
import re
import os
from docx import Document
from pathlib import Path
from docx2pdf import convert as convert_to_pdf
from datetime import datetime
from itertools import zip_longest

class Handover_Pack():
    def __init__(self, cust_num=None):
        #Identify Customer number for the pack
        while True:
            if type(cust_num) == int:
                self.cust_num = str(cust_num)
            elif type(cust_num) == str:
                self.cust_num = cust_num
            else:
                self.cust_num = input("Customer Number: ").strip(" ")
            if len(self.cust_num) < 4:
                print("The Customer number starts with a 4 digits integer. {} is invalid".format(self.cust_num))
                continue
            try:
                int(self.cust_num[:4])
            except ValueError:
                print("The Customer number starts with a 4 digit integer. {} is invalid".format(self.cust_num[:4]))
                continue
            self.paths = backend.get_paths(self.cust_num)
            if self.paths != None:
                print()
                break
            else:
                return None
        self.__init_file_structure__()
        self.__init_values__()
        self.required = {}

        #Establish Completed sections dictionary
        self.check_existing()
        self.section_status()
        self.errors = {}

    def __init_values__(self):
        if self.paths["Pack"].joinpath("Pack Values.txt").exists():
            with open(self.paths["Pack"].joinpath("Pack Values.txt"), "r") as file:
                path_dict = json.load(file)
            self.values = path_dict
        else:
            self.values = {"Business Name":None,
                           "Address":None,
                           "System Size":None,
                           "Predicted Output":None,
                           "Prediction Type":None,
                           "Install Date":None,
                           "Serial Numbers":None}

    def __init_file_structure__(self):
        self.structure = {1:[1.1],
                          2:[2.1],
                          3:[3.1, 3.2, 3.3, 3.4, 3.41, 3.5, 3.6],
                          4:[4.1, 4.2, 4.3, 4.4, 4.5, 4.6],
                          5:[5.1, 5.2],
                          6:[6.1, 6.2],
                          7:[7.1]}
        with open(self.paths["Data"].joinpath("Folder Structure.txt"), "r") as file:
            temp_dict = json.load(file)
        for i in self.structure:
            key = str(i)
            try:
                self.paths[key]
            except:
                self.paths[key] = self.paths["Pack"].joinpath(temp_dict[key])
            self.paths[key].mkdir(exist_ok=True)
            for j in self.structure[i]:
                key2 = str(j)
                try:
                    self.paths[key2]
                except:
                    self.paths[key2] = self.paths[key].joinpath(temp_dict[key2])
                self.paths[key].mkdir(exist_ok=True)
        self.paths["Checklists"] = self.paths["Pack"].joinpath("Checklists")
        self.paths["RunErrors"] = self.paths["Pack"].joinpath("RunErrors")
        self.paths["Archive"] = self.paths["Pack"].joinpath("Archive")
        self.paths["Notes"] = self.paths["Pack"].joinpath("User Notes.txt")
        self.paths["Checklists"].mkdir(exist_ok=True)
        self.paths["RunErrors"].mkdir(exist_ok=True)
        self.paths["Archive"].mkdir(exist_ok=True)
        with open(self.paths["Notes"], 'w') as _:
            pass

    def check_existing(self):
        self.checklist = {1:False, 2:False, 3:False, 4:False, 5:False, 6:False, 7:False,
                          1.1:False, 2.1:False, 3.1:False, 3.2:False, 3.3:False, 3.4:False, 3.41:False,
                          3.5:False, 3.6:False, 4.1:False, 4.2:False, 4.3:False, 4.4:False, 4.5:False,
                          4.6:False, 5.1:False, 5.2:False, 6.1:False, 6.2:False, 7.1:False,
                          }

        for index in self.checklist:
            if index != int(index):
                self.checklist[index] = self.paths[str(index)].exists()
                if not self.checklist[index]:
                    path = backend.open_folder_n(self.paths["Pack"], int(index))
                    for folder in path.iterdir():
                        if index == 3.41:
                            i = "3.4a"
                        else:
                            i = index
                        if folder.parts[-1].startswith("{}.".format(i)):
                            self.checklist[index] = True
                            break
        self.section_status()


    def section_status(self):
        for section in self.structure:
            done = True
            for subsection in self.structure[section]:
                if not self.checklist[subsection]:
                    done = False
            self.checklist[section] = done

    def _require(self, n, add):
        try:
            self.required[n]
        except KeyError:
            self.required[n] = []
        self.required[n].append(add)

    def end_of_run_dump(self):
        dt = datetime.now()
        if self.errors:
            backend.dump_dict(self.paths["RunErrors"].joinpath("RunErrors, {}.txt".format(dt.strftime("%d %b %y, %H-%M-%S"))), self.errors)
        backend.dump_dict(self.paths["Checklists"].joinpath("Checklist, {}.txt".format(dt.strftime("%d %b %y, %H-%M-%S"))), self.checklist)
        #backend.dump_dict(self.paths["Pack"].joinpath("Last Run Checklist.txt"), self.checklist)
        if self.required == {}:
            self.required = {None: "All Information Provided. Pack is complete."}
        backend.dump_dict(self.paths["Pack"].joinpath("Information missing from Last Run.txt"), self.required)
        path_dict = {}
        for key in self.paths:
            if self.paths[key] == None or self.paths[key] == False:
                val = None
            elif type(self.paths[key]) == list:
                val = [str(x) if type(Path()) == type(x) else x for x in self.paths[key]]
            elif type(self.paths[key]) == type(Path()):
                val = str(self.paths[key])
            else:
                val = self.paths[key]
            path_dict[key] = val
        with open(self.paths["Pack"].joinpath("File Paths.txt"), "w") as file:
            json.dump(path_dict, file, indent=3, separators=(',\n', ': '), sort_keys=True)

        val_dict = {}
        for key in self.values:
            if self.values[key] == None or self.values[key] == False:
                val = None
            else:
                val = self.values[key]
            val_dict[key] = val
        with open(self.paths["Pack"].joinpath("Pack Values.txt"), "w") as file:
            json.dump(val_dict, file, indent=3, sort_keys=True)

        with open(self.paths["Main"].joinpath("Pending files.txt"), "w") as file:
            try:
                pending = json.load(file)
            except:
                pending = {}
            if None in self.required:
                self.required.pop(self.paths["Customer"].parts[-1], None)
            else:
                pending[self.paths["Customer"].parts[-1]] = self.required
            json.dump(pending, file, indent=3, sort_keys=True)

    def end_of_run_transfer(self):
        if not self.paths["Customer"] in self.paths["Pack"].parents:
            new_pack_path = backend.open_folder_n(self.paths["Customer"], 11)
            for path in new_pack_path.iterdir():
                arc_path = backend.archive(path, self.paths)
                for key in self.paths:
                    if type(self.paths[key]) == list:
                        self.paths[key] = [x if x != path else arc_path for x in self.paths[key]]
                    else:
                        if self.paths[key] == path:
                            self.paths[key] = arc_path
            for path in self.paths["Pack"].iterdir():
                shutil.move(path, new_pack_path.joinpath(path.parts[-1]))
            for key in self.paths:
                if type(self.paths[key]) == type(Path()):
                    for gen, parent in enumerate(self.paths[key].parents):
                        if parent == self.paths["Pack"]:
                            break
                    else:
                        continue
                    gen = -gen-1
                    new_path = new_pack_path
                    while gen < 0:
                        new_path = new_path.joinpath(self.paths[key].parts[gen])
                        gen += 1
                    self.paths[key] = new_path
                elif type(self.paths[key]) == list:
                    lst = []
                    for p in self.paths[key]:
                        for gen, parent in enumerate(p.parents):
                            if parent == self.paths["Pack"]:
                                break
                        else:
                            continue
                        gen = -gen-1
                        new_path = new_pack_path
                        while gen < 0:
                            new_path = new_path.joinpath(p.parts[gen])
                            gen += 1
                        lst.append(new_path)
                    self.paths[key] = lst
            os.rmdir(self.paths["Pack"])
            self.paths["Pack"] = new_pack_path


    def run(self):
        if self.paths:
            for x in range(1, 8):
                if not self.checklist[x]:
                    exec("self.section_{}()".format(x))
            self.end_of_run_transfer()
            self.end_of_run_dump()

    def section_1(self):
        try:
            if not self.checklist[1.1]:
                try:
                    backend.copy_file(self.paths["Data"].joinpath("Health & Safety Guidelines.pdf"), self.paths["1.1"], overwrite=True)
                    self.checklist[1.1] = True
                except FileNotFoundError:
                    print("Couldn't find \"Health & Safety Guidelines.pdf\" in the Data path.\n")
                    self._require(1, "Missing \"Health & Safety Guidelines.pdf\" in Data folder")
        except:
            print("Error caught in completion of section 1. See RunErrors for details.\n")
            self.errors[1.1] = traceback.format_exc()
        self.section_status()

    def section_2(self):
        try:
            self.paths, self.cust_num, quote_pdf = find.Quotation(self.paths, self.cust_num)
            if not quote_pdf:
                print("Missing the Quotation pdf, cannot complete section 2.")
                self._require(2, "Missing the Quotation pdf")
                return None
            self.paths, self.cust_num, schem_pdf = find.Final_Schematic(self.paths, self.cust_num)

            if not self.checklist[2.1]:
                quote_text = backend.pdf_to_str(quote_pdf)
                if not self.values["Business Name"]:
                    self.values["Business Name"] = backend.find_in_str("Business name", quote_text[0], "\n")
                if not self.values["Address"]:
                    ad = backend.find_in_str("Site Address", quote_text[0], "\n")
                    if ad != re.search("As Above", ad, re.IGNORECASE):
                        ad = backend.find_in_str("Address", quote_text[0], "\n")
                        self.values["Address"] = self.values["Business Name"].strip(".")+", "+ad
                    else:
                        self.values["Address"] = ad
                if not self.values["System Size"]:
                    sys_size = backend.find_in_str("System Size:", quote_text[1], "kWp\n")
                    if sys_size:
                        self.values["System Size"] = float(sys_size)
                    else:
                        os.startfile(self.paths["Quotation"])
                        print("Quotation doesn't have information on the System Size.")
                        self.values["System Size"] = ui.request_float("System Size(kWp)")
                if not self.values["Predicted Output"]:
                    pred_out = backend.find_in_str("estimated generation:", quote_text[1], "kWh\n")
                    if pred_out:
                        self.values["Predicted Output"] = float(pred_out.replace(",",""))
                    else:
                        os.startfile(self.paths["Quotation"])
                        print("Quotation doesn't have information on the Predicted Output.")
                        self.values["Predicted Output"] = ui.request_float("Predicted Output(kWh)")
                if not self.values["Prediction Type"]:
                    os.startfile(str(backend.open_folder_n(self.paths["Customer"], 1)))
                    self.values["Prediction Type"] = ui.choose_from_list(["PVSol", "SolarEdge"], "What type of prediction was used?")
                    print()
                self.paths, self.values = find.Inverter_Information(self.paths, self.values)
                self.paths, self.values = find.Module_Information(self.paths, self.values)
                self.values = find.Install_Date(self.paths, self.values)
                self.values = find.Serial_Numbers(self.paths, self.values)
                if self.values["Serial Numbers"] and self.values["Inverters"]:
                    count, conflicting = ui.check_conflicting_data(len(self.values["Inverters"]), len(self.values["Serial Numbers"]), "the number of inverters")
                    if conflicting:
                        print()
                        print("You have given {} inverters: {}".format(len(self.values["Inverters"]), self.values["Inverters"]))
                        print("and {} serial numbers: {}".format(len(self.values["Serial Numbers"]), self.values["Serial Numbers"]))
                        if count == len(self.values["Serial Numbers"]):
                            print("As you indicated there should be {} inverters, the inverter information has been cleared.".format(count))
                            print("This section will now abort and the assistor will need to be run again.\n")
                            self.values["Inverters"] = None
                            self.values["SolarEdge Warranty"] = None
                            self.paths["Inverter Datasheets"] = None
                            return None
                        elif count == len(self.values["Inverters"]):
                            print("As you indicated there should be {} inverters, the serial number information has been cleared.".format(count))
                            print("This section will now abort and the assistor will need to be run again.\n")
                            self.values["Serial Numbers"] = None
                            return None
                        elif count == None:
                            print("As you indicated uncertainty, both have been cleared to prevent other sections completing falsely")
                            print("This section will now abort and the assistor will need to be run again.\n")
                            self.values["Inverters"] = None
                            self.values["SolarEdge Warranty"] = None
                            self.paths["Inverter Datasheets"] = None
                            self.values["Serial Numbers"] = None
                            return None

                checklist = {"Address": False, "Business Name": False, "Install Date": False, "Inverters": False,
                         "Module": False, "Module Number": False, "Predicted Output": False, "Prediction Type": False,
                         "Serial Numbers": False, "SolarEdge Warranty": False, "System Size": False}
                check = True
                for key in checklist:
                    if not self.values[key]:
                        check = False
                    else:
                        checklist[key] = True

                if check:
                    path = self.paths["2.1"].with_suffix(".docx")
                    try:
                        backend.copy_file(self.paths["Data"].joinpath("Information Template.docx"), path, overwrite=True)
                    except FileNotFoundError:
                        print("Couldn't find 'Information Template.docx' in the Data path.\nSkipping Section 2\n")
                        return None
                    document = Document(path)
                    document.paragraphs[4].runs[0].text = "{:.2f} kWp".format(self.values["System Size"])
                    document.tables[1].rows[1].cells[1].paragraphs[0].runs[0].text = "{:.2f}".format(self.values["System Size"])
                    document.tables[1].rows[2].cells[1].paragraphs[0].runs[0].text = "{:,.0f} kW".format(self.values["Predicted Output"])
                    document.tables[1].rows[2].cells[1].paragraphs[0].runs[5].text = ""
                    document.tables[1].rows[2].cells[1].paragraphs[0].runs[7].text = ""
                    document.tables[1].rows[2].cells[1].paragraphs[0].runs[8].text = ""
                    if self.values["Prediction Type"] == "PVSol":
                        document.tables[1].rows[2].cells[1].paragraphs[0].runs[6].text = ""
                    elif self.values["Prediction Type"] == "SolarEdge":
                        document.tables[1].rows[2].cells[1].paragraphs[0].runs[4].text = ""
                    inverters = [[x,self.values["Inverters"].count(x)] for x in set(self.values["Inverters"])]
                    inverter_str = ""
                    for inv in inverters:
                        inverter_str += " & "+inv[0]
                        if inv[1] != 1:
                            inverter_str += " x {}".format(inv[1])
                    inverter_str = inverter_str.strip(" & ")
                    document.tables[1].rows[4].cells[1].paragraphs[0].runs[0].text = inverter_str
                    document.tables[1].rows[3].cells[1].paragraphs[0].runs[0].text = self.values["Module"]+" x "+str(self.values["Module Number"])

                    #Format Address in Box
                    lst = self.values["Address"].split(",")
                    text = []
                    lines = 0
                    for a,b in zip_longest(lst[::2],lst[1::2]):
                        lines += 1
                        if b:
                            text.append(a+","+b+",\n")
                        else:
                            text.append(a+",\n")
                    lst = text
                    text = lst[0]
                    for st in lst[1:]:
                        text += st
                    text = text.strip(",\n")
                    document.tables[0].rows[1].cells[0].paragraphs[3].runs[0].text = text
                    for i in range(lines-1):
                        p = document.tables[0].rows[1].cells[0].paragraphs[4]._element
                        p.getparent().remove(p)
                        p._p = p._element = None
                    document.tables[0].rows[1].cells[0].paragraphs[8-lines].runs[0].text = "Job ref: "+self.cust_num
                    #End of Address formatting

                    if len(self.values["Serial Numbers"]) > 4:
                        print("There are too many serial numbers to automatically format into the table for the \"2.1  System Summary & General Information\". Please format manually when prompted before saving.")
                    else:
                        for i, sn in enumerate(self.values["Serial Numbers"]):
                            document.tables[1].rows[5+i].cells[1].paragraphs[0].runs[0].text = sn
                    document.tables[1].rows[9].cells[1].paragraphs[0].runs[0].text = self.values["Install Date"]

                    document.save(path)
                    os.startfile(str(path))
                    while True:
                        confirm = input("Confirm Document formating, then save and exit, then enter \"Done\". Enter \"Skip\" if formating can't be resolved.\n")
                        if confirm == "Done":
                            to_pdf = True
                            break
                        elif confirm == "Skip":
                            to_pdf = False
                            break
                        else:
                            print("Please enter \"Done\" or \"Skip\"")
                    if to_pdf:
                        convert_to_pdf(path)
                        self.checklist[2.1] = True
                    else:
                        self._require(2, "Didn't complete formatting of Word Document before conversion to pdf.")

                    backend.archive(path, self.paths)
                else:
                    for key in checklist:
                        if not checklist[key]:
                            self._require(2, "Missing value for \"{}\"".format(key))
        except:
            print("Error caught in completion of section 2. See RunErrors for details.\n")
            self.errors[2] = traceback.format_exc()
        self.section_status()

    def section_3(self):
        try:
            if not self.checklist[3.1]:
                self.paths = ui.request_warranty_path(self.paths)
                if self.paths["Warranty"]:
                    backend.copy_file(self.paths["Warranty"], self.paths["3.1"], overwrite=True)
                else:
                    self._require(3, "Missing Mypower Installation Warranty.")
                self.checklist[3.1] = True

            if not self.checklist[3.2] or not self.checklist[3.3]:
                self.paths, self.values = find.Module_Information(self.paths, self.values)
                if self.paths["Module Warranty"]:
                    backend.copy_file(self.paths["Module Warranty"], self.paths["3.2"], overwrite=True)
                if self.paths["Module Datasheet"]:
                    backend.copy_file(self.paths["Module Datasheet"], self.paths["3.3"], overwrite=True)
                if not (self.paths["Module Warranty"] and self.paths["Module Datasheet"]):
                    self._require(3, "Missing Module type for warranties and datasheets.")

            if not self.checklist[3.5]:
                self.paths, self.values = find.Inverter_Information(self.paths, self.values)
                if self.values["SolarEdge Warranty"]:
                    self.paths = ui.request_solaredge_warranty_path(self.paths)
                    if self.paths["SolarEdge Warranty"]:
                        backend.copy_file(self.paths["SolarEdge Warranty"], self.paths["3.5"], overwrite=True)
                        self.checklist[3.5] = True
                elif self.values["SolarEdge Warranty"] == False:
                    self.checklist[3.5] = True
                else:
                    self._require(3, "Is a SolarEdge Warranty required?")

            if not self.checklist[3.4]:
                self.paths, self.values = find.Inverter_Information(self.paths, self.values)
                if self.paths["Inverter Datasheets"] and self.values["Inverters"]:
                    if len(self.paths["Inverter Datasheets"]) == 1 or all(x==self.paths["Inverter Datasheets"][0] for x in self.paths["Inverter Datasheets"]):
                        backend.copy_file(self.paths["Inverter Datasheets"][0], self.paths["3.4"], overwrite=True)
                    else:
                        groups = []
                        for name, datasheet in zip(self.values["Inverters"], self.paths["Inverter Datasheets"]):
                            if (name, datasheet) not in groups:
                                groups.append((name, datasheet))
                        for i, (name, datasheet) in enumerate(groups):
                            path = self.paths["3.4"].parent.joinpath(self.paths["3.4"].with_suffix("").parts[-1]+" ({}).pdf".format(name))
                            path = path.parent.joinpath("3.4.{}".format(i+1)+path.parts[-1].strip("3.4"))
                            backend.copy_file(datasheet, path, overwrite=True)
                    self.checklist[3.4] = True
                else:
                    self._require(3, "Missing Inverter type for warranties and datasheets.")

            if not self.checklist[3.6]:
                self.paths, self.values = find.Optimiser_Information(self.paths, self.values)
                if self.paths["Optimiser Datasheet"]:
                    if not self.values["Optimiser"] == "No Optimisers":
                        backend.copy_file(self.paths["Optimiser Datasheet"], self.paths["3.6"], overwrite=True)
                    self.checklist[3.6] = True
                else:
                    self._require(3, "Missing Optimiser type for warranties and datasheets.")

            if not self.checklist[3.41]:
                self.paths, self.values = find.Inverter_Information(self.paths, self.values)
                self.values = find.Serial_Numbers(self.paths, self.values)
                self.paths, self.values = find.Extended_Warranties(self.paths, self.values)
                if not any([True if x == None else False for x in self.paths["Extended Warranties"]]):
                    if len(self.paths["Extended Warranties"]) == 1:
                        confirm = backend.copy_file(self.paths["Extended Warranties"][0], self.paths["3.41"], overwrite=False)
                        if not confirm:
                            print("File already exists at the location the Extended Warranty was trying to be moved to. Please move this to the archive and try again.")
                            raise FileExistsError("File already exists at the location the Extended Warranty was trying to be moved to. Please move this to the archive and try again.")
                    else:
                        groups = []
                        for warr,sn in zip(self.paths["Extended Warranties"], self.values["Serial Numbers"]):
                            if (warr, sn) not in groups and warr:
                                groups.append((warr, sn))
                        for i, (warr, sn) in enumerate(groups):
                            path = self.paths["3.41"].parent.joinpath(self.paths["3.41"].with_suffix("").parts[-1]+" ({}).pdf".format(sn))
                            path = path.parent.joinpath("3.4a.{}".format(i+1)+path.parts[-1].strip("3.4a"))
                            confirm = backend.copy_file(warr, path, overwrite=False)
                            if not confirm:
                                print("File already exists at the location the Extended Warranty was trying to be moved to. Please move this to the archive and try again.")
                                raise FileExistsError("File already exists at the location the Extended Warranty was trying to be moved to. Please move this to the archive and try again.")
                    self.checklist[3.41] = True
                else:
                    self._require(3, "Missing Extended Warranties.")
        except:
            print("Error caught in completion of section 3. See RunErrors for details.\n")
            self.errors[3] = traceback.format_exc()
        self.section_status()

    def section_4(self):
        try:
            if not self.checklist[4.1]:
                self.paths, self.cust_num, schem_pdf = find.Final_Schematic(self.paths, self.cust_num)
                if schem_pdf:
                    backend.copy_file(schem_pdf, self.paths["4.1"])
                    self.checklist[4.1] = True
                else:
                    self._require(4, "Missing Final Schematic.")
            if not self.checklist[4.2]:
                self.paths["AC Cert"] = ui.find_in_folder(self.paths, "Commissioning test report (AC Cert)", n=10)
                if self.paths["AC Cert"]:
                    backend.copy_file(self.paths["AC Cert"], self.paths["4.2"])
                    self.checklist[4.2] = True
                else:
                    self._require(4, "Missing Commissioning test reports (AC Cert).")
            if not self.checklist[4.3]:
                self.paths["DC Cert"] = ui.find_in_folder(self.paths, "Commissioning test report (DC Cert)", n=10)
                if self.paths["DC Cert"]:
                    backend.copy_file(self.paths["DC Cert"], self.paths["4.3"])
                    self.checklist[4.3] = True
                else:
                    self._require(4, "Missing Commissioning test reports (DC Cert).")
            if not self.checklist[4.4]:
                self.paths["G99 Form"] = ui.find_in_folder(self.paths, "DNO commissioning form (G99 Form A3-1)", n=6)
                if self.paths["G99 Form"]:
                    backend.copy_file(self.paths["G99 Form"], self.paths["4.4"])
                    self.checklist[4.4] = True
                else:
                    self._require(4, "Missing DNO commissioning form (G99 Form A3-1).")
            if not self.checklist[4.5]:
                self.paths["Sign Off"] = ui.find_in_folder(self.paths, "Inverter & wiring sign off", n=10)
                if self.paths["Sign Off"]:
                    backend.copy_file(self.paths["Sign Off"], self.paths["4.5"])
                    self.checklist[4.5] = True
                else:
                    self._require(4, "Missing Inverter & wiring sign off.")
            if not self.checklist[4.6]:
                self.paths["DNO Notification"] = ui.find_in_folder(self.paths, "DNO commissioning notification", n=6)
                if self.paths["DNO Notification"]:
                    backend.copy_file(self.paths["DNO Notification"], self.paths["4.6"])
                    self.checklist[4.6] = True
                else:
                    self._require(4, "Missing DNO commissioning notification.")
        except:
            print("Error caught in completion of section 4. See RunErrors for details.\n")
            self.errors[4] = traceback.format_exc()
        self.section_status()

    def section_5(self):
        try:
            if not self.checklist[5.1]:
                self.paths["Summary Report"] = ui.find_in_folder(self.paths, "Summary Report", n=1)
                if self.paths["Summary Report"]:
                    backend.copy_file(self.paths["Summary Report"], self.paths["5.1"])
                    self.checklist[5.1] = True
                else:
                    self._require(5, "Missing Summary Report.")
            if not self.checklist[5.2]:
                self.paths["Prediction Tool"] = ui.find_in_folder(self.paths, "Predicted Output Comparison Tool", n=11, file_types=[".xlsx"])
                if self.paths["Prediction Tool"]:
                    backend.copy_file(self.paths["Prediction Tool"], self.paths["5.2"])
                    self.checklist[5.2] = True
                else:
                    self._require(5, "Missing Predicted Output Comparison Tool.")
        except:
            print("Error caught in completion of section 5. See RunErrors for details.\n")
            self.errors[5] = traceback.format_exc()
        self.section_status()

    def section_6(self):
        try:
            if not self.checklist[6.1]:
                self.paths["NAPIT Notification"] = ui.find_in_folder(self.paths, "NAPIT Work notification details", n=11)
                if self.paths["NAPIT Notification"]:
                    backend.copy_file(self.paths["NAPIT Notification"], self.paths["6.1"])
                    self.checklist[6.1] = True
                else:
                    self._require(6, "Missing NAPIT Work notification details.")
            if not self.checklist[6.2]:
                try:
                    self.paths["Structural Cert"]
                except KeyError:
                    self.paths["Structural Cert"] = None
                if self.paths["Structural Cert"] == None:
                    self.paths["Structural Cert"] = ui.find_in_folder(self.paths, "Structural survey certificate", option_for_none=True, n=1)
                if self.paths["Structural Cert"]:
                    if self.paths["Structural Cert"] != True:
                        backend.copy_file(self.paths["Structural Cert"], self.paths["6.2"])
                    self.checklist[6.2] = True
                elif self.paths["Structural Cert"] == False:
                    self.checklist[6.2] = True
                else:
                    self._require(6, "Missing Structural survey certificate.")
        except:
            print("Error caught in completion of section 6. See RunErrors for details.\n")
            self.errors[6] = traceback.format_exc()
        self.section_status()

    def section_7(self):
        try:
            if not self.checklist[7.1]:
                if not self.values["System Size"]:
                    self.paths, self.cust_num, quote_pdf = find.Quotation(self.paths, self.cust_num)
                    if quote_pdf:
                        self.paths, self.cust_num, schem_pdf = find.Final_Schematic(self.paths, self.cust_num)
                        quote_text = backend.pdf_to_str(quote_pdf)
                        sys_size = backend.find_in_str("System Size:", quote_text[1], "kWp\n")
                        if sys_size:
                            self.values["System Size"] = float(sys_size)
                if self.values["System Size"]==None:
                    try:
                        self.paths["MCS Certificate"]
                    except KeyError:
                        self.paths["MCS Certificate"] = None
                    if self.paths["MCS Certificate"] == None:
                        self.paths["MCS Certificate"] = ui.find_in_folder(self.paths, "MCS Certificate", n=11, option_for_none=True)
                    if self.paths["MCS Certificate"]:
                        if self.paths["MCS Certificate"] != True:
                            backend.copy_file(self.paths["MCS Certificate"], self.paths["7.1"])
                        self.checklist[7.1] = True
                    else:
                        self._require(7, "Missing MCS Certificate.")
                elif self.values["System Size"]<=50:
                    self.paths["MCS Certificate"] = ui.find_in_folder(self.paths, "MCS Certificate", n=11)
                    if self.paths["MCS Certificate"]:
                        backend.copy_file(self.paths["MCS Certificate"], self.paths["7.1"])
                        self.checklist[7.1] = True
                    else:
                        self._require(7, "Missing MCS Certificate.")
                if self.values["System Size"]>50 or self.paths["MCS Certificate"]==False:
                    os.rmdir(self.paths["7"])
                    self.checklist[7.1] = True
        except:
            print("Error caught in completion of section 7. See RunErrors for details.\n")
            self.errors[7] = traceback.format_exc()
        self.section_status()
